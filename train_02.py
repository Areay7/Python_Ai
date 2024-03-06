# 第1题(代码):
# 导入 math 模块
import math
# 提示用户输入半径，并将输入的字符串存储在变量 r_str 中
r_str = input('请输入一个半径：')
# 检查用户输入是否是有效的数字（包括小数），通过去掉第一个点并检查是否全是数字来判断
if r_str.replace('.', '', 1).isdigit():
    # 将字符串 r_str 转换为浮点数，并将其赋值给变量 r
    r = float(r_str)
    # 检查半径是否大于0
    if r > 0:
        # 使用公式 π * r^2 计算圆的面积，并将结果赋值给变量 s
        s = math.pi * r ** 2
        # 打印圆的面积
        print("圆的面积为:", s)
    else:
        # 如果半径不大于0，打印一条消息指示半径必须大于0
        print('半径必须大于0。')
else:
    # 如果用户输入不是有效的数字，打印一条消息指示输入无效
    print('输入无效，请确保输入的是数字。')

# 第2题(代码):
# 提示用户输入总人数
n = input("请输入总人数：")
# 将输入的字符串转换为整数，并赋值给变量 n
n = int(n)
# 提示用户输入报到第几个数退出
k = input("报到第几个数退出：")
# 将输入的字符串转换为整数，并赋值给变量 k
k = int(k)
# 创建一个空列表 alist，用于存储初始编号为1到n的人
alist = []
for i in range(1, n + 1):
    alist.append(i)
# 打印初始列表
print(alist)
# 初始化变量 t 和 i
t = 1
i = 1
# 使用循环模拟报数并删除过程，直到只剩下一个人
while n > 1:
    # 计算要删除的位置
    t = t + k - 1
    # 如果计算结果超过当前列表长度，取余数
    if t > n:
        t = t % n
    # 打印每次删除的人的编号
    print("第", i, "次删除：", alist[t - 1])
    # 更新循环计数器
    i = i + 1
    # 删除列表中对应位置的元素
    del alist[t - 1]
    # 更新剩余人数
    n = n - 1
# 打印最后留下的人的编号
print("最后留下的是原来的第", alist[0], "号")

# 第3题(代码):
from docx import Document
d = Document('ttt.docx')
for p in d.paragraphs:
 for index, run in enumerate(p.runs):
    if run.style.name == 'Hyperlink':
       print(run.text,end=':')
       for child in p.runs[index-2].element.getchildren():
        text = child.text
        if text and text.startswith('HYPERLINK'):
           print(text[:])


# 第4题(代码):
# 导入 Document 类和 RGBColor 类
from docx import Document
from docx.shared import RGBColor
# 初始化列表，用于存储加粗字体和红色字体的文本
boldText = []
redText = []
# 使用 Document 类加载指定的 Word 文档
doc = Document('ttt.docx')
# 遍历文档中的每个段落
for p in doc.paragraphs:
    # 遍历每个段落中的每个运行（run）
    for r in p.runs:
        # 检查是否为加粗字体，如果是，将文本添加到 boldText 列表
        if r.bold:
            boldText.append(r.text)
        # 检查字体颜色是否为红色，如果是，将文本添加到 redText 列表
        if r.font.color.rgb == RGBColor(255, 0, 0):
            redText.append(r.text)
# 构建结果字典，包括红色字体文本、加粗字体文本和两者的交集
result = {'red text': redText,
          'bold text': boldText,
          'both': set(redText) & set(boldText)}
# 遍历结果字典，打印输出
for title in result.keys():
    print(title.center(30, '='))
    for text in result[title]:
        print(text)

